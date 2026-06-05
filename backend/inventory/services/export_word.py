from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(document, text, level=1):
    heading = document.add_heading(text, level=level)
    return heading


def add_currency_text(value):
    return f"{int(value):,}đ".replace(",", ".")


def create_monthly_report_document(
    month,
    year,
    total_revenue,
    invoice_count,
    sold_quantity,
    top_products,
    low_stock_variants,
):
    document = Document()

    title = document.add_heading("BÁO CÁO KINH DOANH THÁNG", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(f"Tháng {month}/{year}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    add_heading(document, "1. Tổng quan", level=1)

    overview_table = document.add_table(rows=1, cols=2)
    overview_table.style = "Table Grid"

    header_cells = overview_table.rows[0].cells
    header_cells[0].text = "Chỉ tiêu"
    header_cells[1].text = "Giá trị"

    rows = [
        ("Tổng doanh thu", add_currency_text(total_revenue)),
        ("Số hóa đơn", str(invoice_count)),
        ("Số sản phẩm bán ra", str(sold_quantity)),
    ]

    for label, value in rows:
        row_cells = overview_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    document.add_paragraph("")

    add_heading(document, "2. Top sản phẩm bán chạy", level=1)

    if top_products:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        cells = table.rows[0].cells
        cells[0].text = "Mã hàng"
        cells[1].text = "Tên sản phẩm"
        cells[2].text = "Số lượng bán"
        cells[3].text = "Doanh thu"

        for item in top_products:
            row = table.add_row().cells
            row[0].text = item["product_variant__product__code"]
            row[1].text = item["product_variant__product__name"]
            row[2].text = str(item["total_quantity"] or 0)
            row[3].text = add_currency_text(item["total_revenue"] or 0)
    else:
        document.add_paragraph("Chưa có dữ liệu bán hàng trong tháng này.")

    document.add_paragraph("")

    add_heading(document, "3. Hàng sắp hết", level=1)

    if low_stock_variants:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        cells = table.rows[0].cells
        cells[0].text = "Mã hàng"
        cells[1].text = "Tên sản phẩm"
        cells[2].text = "Size"
        cells[3].text = "Màu"
        cells[4].text = "Tồn kho"

        for variant in low_stock_variants:
            row = table.add_row().cells
            row[0].text = variant.product.code
            row[1].text = variant.product.name
            row[2].text = variant.size
            row[3].text = variant.color
            row[4].text = f"{variant.current_stock}/{variant.low_stock_threshold}"
    else:
        document.add_paragraph("Không có mặt hàng nào dưới ngưỡng cảnh báo tồn kho.")

    document.add_paragraph("")

    add_heading(document, "4. Nhận xét tổng quan", level=1)

    if total_revenue > 0:
        document.add_paragraph(
            f"Trong tháng {month}/{year}, đại lý ghi nhận tổng doanh thu "
            f"{add_currency_text(total_revenue)} từ {invoice_count} hóa đơn, "
            f"với tổng số {sold_quantity} sản phẩm đã bán ra."
        )
    else:
        document.add_paragraph(
            f"Trong tháng {month}/{year}, hệ thống chưa ghi nhận doanh thu bán hàng."
        )

    if low_stock_variants:
        document.add_paragraph(
            f"Có {len(low_stock_variants)} mặt hàng đang ở mức tồn kho thấp. "
            f"Đại lý nên kiểm tra và cân nhắc nhập bổ sung để tránh thiếu hàng."
        )
    else:
        document.add_paragraph(
            "Tình trạng tồn kho hiện tại ổn định, chưa có mặt hàng nào cần cảnh báo nhập thêm."
        )

    document.add_paragraph("")
    document.add_paragraph("Báo cáo được tạo tự động từ hệ thống Viettien Agency Manager.")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)

    return document